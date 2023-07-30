import json, os, datetime, urllib, requests, logging, re, csv, smtplib, slack, psycopg2, base64, hashlib, string
from Crypto.Cipher import AES
from Crypto import Random
from requests.packages.urllib3.exceptions import InsecureRequestWarning
from docx import Document
from jira.client import JIRA
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

requests.packages.urllib3.disable_warnings(InsecureRequestWarning)
Current_User_Agent: str = 'Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:78.0) Gecko/20100101 Firefox/78.0'

def Alnum_List() -> list:
    Lowercase: list = list(string.ascii_lowercase)
    Uppercase: list = list(string.ascii_uppercase)
    Numbers: list = [str(r) for r in range(0, 10)]
    AlNum: list = Lowercase.copy()
    AlNum.extend(Uppercase)
    AlNum.extend(Numbers)
    return AlNum

class Fang:

    def __init__(self):
        self.entity_components: dict = {
            ".": "[.]",
            "http": "hXXp",
            "@": "[AT]",
            "ftp": "fXp",
            "magnet": "mXgnXt"
        }

    def Defang(self, entity: str = str()) -> str:
        
        for fanged, defanged in self.entity_components.items():
            entity = entity.replace(fanged, defanged)

        return entity

    def Fang(self, entity: str = str()) -> str:

        for fanged, defanged in self.entity_components.items():
            entity = entity.replace(defanged, fanged)

        return entity

    def Defang_List(self, entity: list = list()) -> list:
        iterator: int = int()

        for item in entity:
        
            for fanged, defanged in self.entity_components.items():
                item = item.replace(fanged, defanged)

            entity[iterator] = item
            iterator += 1

        return entity

    def Fang_List(self, entity: list = list()) -> list:
        iterator: int = int()

        for item in entity:
        
            for fanged, defanged in self.entity_components.items():
                item = item.replace(defanged, fanged)

            entity[iterator] = item
            iterator += 1

        return entity

class Coder:

    def __init__(self, to_code):
        self.to_code = to_code

    def encode_check(self):

        if isinstance(self.to_code, bytes):
            self.to_code = self.to_code

        else:
            self.to_code = self.to_code.encode()

    def decode_check(self):

        if isinstance(self.to_code, str):
            self.to_code = self.to_code

        else:
            self.to_code = self.to_code.decode()

    def b64_encode(self):
        self.encode_check()
        return base64.b64encode(self.to_code).decode()

    def b64_urlsafe_encode(self):
        self.encode_check()
        return base64.urlsafe_b64encode(self.to_code).decode()

    def b64_decode(self, decode: bool=True):
        self.decode_check()

        if decode:
            return base64.b64decode(self.to_code).decode()

        else:
            return base64.b64decode(self.to_code)

    def b64_urlsafe_decode(self, decode):
        self.decode_check()

        if decode:
            return base64.urlsafe_b64decode(self.to_code).decode()

        else:
            return base64.urlsafe_b64decode(self.to_code)

class Cryptography:

    def __init__(self):
        """Initialises cryptography object"""
        BLOCK_SIZE: int = 16
        self.pad = lambda s: s + (BLOCK_SIZE - len(s) % BLOCK_SIZE) * chr(BLOCK_SIZE - len(s) % BLOCK_SIZE)
        self.unpad = lambda s: s[:-ord(s[len(s) - 1:])]
        self.filesystem_uuid = os.environ["DISK_UUID"]

        if self.filesystem_uuid is None:
            Message = "Environment Variables needed for cryptography don't exist."
            logging.warning(Message)
            raise ValueError(Message)

    def encrypt(self, raw):
        """Encrypts data"""
        private_key = hashlib.sha256(self.filesystem_uuid.encode("utf-8")).digest()
        raw = self.pad(raw)
        iv = Random.new().read(AES.block_size)
        cipher = AES.new(private_key, AES.MODE_CBC, iv)
        return Coder(iv + cipher.encrypt(raw.encode())).b64_encode()

    def decrypt(self, enc):
        """Decrypts data"""
        private_key = hashlib.sha256(self.filesystem_uuid.encode("utf-8")).digest()
        enc = Coder(enc).b64_decode(decode=False)
        iv = enc[:16]
        cipher = AES.new(private_key, AES.MODE_CBC, iv)
        return self.unpad(cipher.decrypt(enc[16:])).decode()

    def configuration_encrypt(self, configuration: dict() = dict()) -> str:
        """Encrypts configuration data"""
        return self.encrypt(JSON_Handler(configuration).Dump_JSON(Sort=False))

    def configuration_decrypt(self, encrypted_configuration: str = str()) -> dict:
        """Decrypts configuration data"""
        # Requires double JSON loading.
        return JSON_Handler(JSON_Handler(self.decrypt(encrypted_configuration)).To_JSON_Loads()).To_JSON_Loads()

class Configuration:

    def __init__(self, Input: bool = bool(), Output: bool = bool(), Core: bool = bool()):
        """Initialises configuration class"""

        try:
            self.key: str = str()

            if Input:
                self.key: str = "inputs"

            elif Output:
                self.key: str = "outputs"

            elif Core:
                self.key: str = "core"

            if self.key != str():
                JSON_File = open(Set_Configuration_File(), "r")
                Configuration_Data = Cryptography().configuration_decrypt(JSON_File.read())
                JSON_File.close()
                self.JSON_Data = Configuration_Data
                self.JSON_Object_Details = Configuration_Data[self.key]

            else:
                self.JSON_Object_Details = None

        except Exception as e:
            logging.warning(f"{Date()} - Common Library - Failed to load configuration file - {str(e)}.")

    def Load_Keys(self):
        """Loads only the configuration keys"""

        try:

            if self.JSON_Object_Details:
                return list(self.JSON_Object_Details.keys())

            else:
                return None

        except Exception as e:
            logging.warning(f"{Date()} - Common Library - Failed to load keys for section {self.key} - {str(e)}.")

    def Load_Values(self, Object: str= str()):
        """Loads only the configuration values"""

        try:

            if self.JSON_Object_Details and Object != str():
                return self.JSON_Object_Details[Object]

            else:
                return None

        except Exception as e:
            logging.warning(f"{Date()} - Common Library - Failed to load object data - {str(e)}.")

    def Set_Field(self, Object: str= str(), Config: dict = dict()):
        """Updates configuration values for a given object"""

        try:

            if self.JSON_Object_Details and Object != str() and Config != dict() and Object in self.JSON_Data[self.key]:
                self.JSON_Data[self.key][Object] = Config
                JSON_File = open(Set_Configuration_File(), "w")
                JSON_File.write(Cryptography().configuration_encrypt(JSON_Handler(self.JSON_Data).Dump_JSON()))
                JSON_File.close()
                return True

        except Exception as e:
            logging.warning(f"{Date()} - Common Library - Failed to set field - {str(e)}.")  

    def Load_Configuration(self, Location: bool = bool(), Postgres_Database: bool = bool(), Object: str = str(), Details_to_Load: list = list()):
        """Loads configuration details for plugins to interact with third party sites, and for outputs to export results"""

        try:

            if Postgres_Database:
                Details_to_Load = ["host", "port", "user", "password", "database"]

            if self.JSON_Object_Details:
                
                if Object in self.JSON_Object_Details:
                    Return_Details: list = list()

                    for Detail in Details_to_Load:
                        Current_Item = self.JSON_Object_Details[Object]

                        if Detail in Current_Item:
                            Current_Item = Current_Item[Detail]

                            if type(Current_Item) in [bool, int]:
                                Return_Details.append(Current_Item)

                            elif Current_Item:
                                Return_Details.append(Current_Item)

                            else:
                                return None

                        else:
                            return None

                    if not Location and not Postgres_Database:

                        if len(Return_Details) > 1:
                            return Return_Details

                        elif len(Return_Details) == 1:
                            return Return_Details[0]

                        else:
                            return None

                    elif Location and not Postgres_Database:
                        Result_Detail = Return_Details[0]
                        Valid_Locations: tuple = (
                            'ac', 'ad', 'ae', 'af', 'ag', 'ai', 'al', 'am', 'ao', 'aq', 'ar', 'as', 'at', 'au', 'az', 'ba', 'bd', 'be', 'bf', 'bg', 'bh', 'bi', 'bj', 'bn', 'bo', 'br', 'bs', 'bt', 'bw', 'by', 'bz', 'ca', 'cc', 'cd', 'cf',
                            'cg', 'ch', 'ci', 'ck', 'cl', 'cm', 'cn', 'co', 'cr', 'cu', 'cv', 'cy', 'cz', 'de', 'dj', 'dk', 'dm', 'do', 'dz', 'ec', 'ee', 'eg', 'es', 'et', 'eu', 'fi', 'fj', 'fm', 'fr', 'ga', 'ge', 'gf', 'gg', 'gh', 'gi',
                            'gl', 'gm', 'gp', 'gr', 'gt', 'gy', 'hk', 'hn', 'hr', 'ht', 'hu', 'id', 'ie', 'il', 'im', 'in', 'io', 'iq', 'is', 'it', 'je', 'jm', 'jo', 'jp', 'ke', 'kg', 'kh', 'ki', 'kr', 'kw', 'kz', 'la', 'lb', 'lc', 'li',
                            'lk', 'ls', 'lt', 'lu', 'lv', 'ly', 'ma', 'md', 'me', 'mg', 'mk', 'ml', 'mm', 'mn', 'ms', 'mt', 'mu', 'mv', 'mw', 'mx', 'my', 'mz', 'na', 'ne', 'nf', 'ng', 'ni', 'nl', 'no', 'np', 'nr', 'nu', 'nz', 'om', 'pa',
                            'pe', 'pf', 'pg', 'ph', 'pk', 'pl', 'pn', 'pr', 'ps', 'pt', 'py', 'qa', 're', 'ro', 'rs', 'ru', 'rw', 'sa', 'sb', 'sc', 'se', 'sg', 'sh', 'si', 'sk', 'sl', 'sm', 'sn', 'so', 'sr', 'st', 'sv', 'sy', 'td', 'tg',
                            'th', 'tj', 'tk', 'tl', 'tm', 'tn', 'to', 'tt', 'tz', 'ua', 'ug', 'uk', 'us', 'uy', 'uz', 'vc', 've', 'vg', 'vi', 'vn', 'vu', 'ws', 'za', 'zm', 'zw'
                        )

                        if Result_Detail not in Valid_Locations:
                            logging.warning(f"{Date()} - Common Library - An invalid location has been specified, please provide a valid location in the config.json file.")

                        else:
                            logging.info(f"{Date()} - Common Library - Country code {Result_Detail} selected.")
                            return Result_Detail

                    else:

                        try:
                            DB_Connection = psycopg2.connect(user=Return_Details[2],
                                          password=Return_Details[3],
                                          host=Return_Details[0],
                                          port=int(Return_Details[1]),
                                          database=Return_Details[4])

                            if DB_Connection:
                                return DB_Connection

                            else:
                                return None

                        except Exception as e:
                            logging.warning(f"{Date()} - Common Library - Failed to connect to database - {str(e)}.")
                            return None

                else:
                    return None

            else:
                return None

        except Exception as e:
            logging.warning(f"{Date()} - Common Library - Failed to load details for object {str(Object)} - {str(e)}.")

def CSV_Output(Object, Title, Plugin_Name, Domain, Link, Result_Type, Output_File, Task_ID, Directory):
    """"""

    try:
        Use_CSV = Load_Output(Object, "csv")

        if Use_CSV:
            File_Dir: str = os.path.dirname(os.path.realpath('__file__'))
            Headings: tuple = ("Title", "Plugin", "Domain", "Link", "Created At", "Output Files", "Result Type", "Task ID")
            Data: tuple = (Title, Plugin_Name, Domain, Link, Date(), Output_File, Result_Type, str(Task_ID))
            Complete_File = f"{File_Dir}/static/protected/output/{Directory}/{Plugin_Name}-Output.csv"

            if not os.path.exists(Complete_File):
                CSV_Output = csv.writer(open(Complete_File, 'w'))
                CSV_Output.writerow(Headings)
                CSV_Output.writerow(Data)
                logging.info(f"{Date()} - Common Library - Created new CSV file located at {str(Complete_File)}.")

            else:
                CSV_Output = csv.writer(open(Complete_File, 'a'))
                CSV_Output.writerow(Data)
                logging.info(f"{Date()} - Common Library - Updated existing CSV file located at {str(Complete_File)}.")

            return Complete_File

        else:
            return None

    except Exception as e:
        logging.warning(f"{Date()} - Common Library - {str(e)}.")

def DOCX_Output(Object, Title, Plugin_Name, Domain, Link, Result_Type, Output_File, Task_ID, Directory):

    try:
        Use_DOCX = Load_Output(Object, "docx")

        if Use_DOCX:
            print(1)
            File_Dir: str = os.path.dirname(os.path.realpath('__file__'))
            Template_File: str = f"{File_Dir}/plugins/common/templates/Scrummage_Report_Template.docx"
            Output_File: str = f"{File_Dir}/static/protected/output/{Directory}/{Plugin_Name}-Output.docx"
            print(2)

            if os.path.exists(Output_File):
                document = Document(Output_File)
                Finding_Style = document.styles['Finding Sub Heading']

            else:
                print(3)
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                Document_Title = f"Scrummage Finding Report for the {Plugin_Name} Plugin"
                document = Document(Template_File)
                Title_Style = document.styles['Title']
                Finding_Style = document.styles['Finding Sub Heading']
                p = document.add_paragraph(Document_Title, style=Title_Style)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                document.add_page_break()
                header = document.sections[0].header
                h = header.paragraphs[0]
                h.text = Document_Title
                print(4)
                h.style = document.styles["Normal"]
                document.add_heading("About", 1)
                document.add_paragraph("Scrummage is an Open-Source Intelligence (OSINT) gathering tool. It helps individuals and organisations alike to measure their online security posture by correlating information offered by various, third-party, services. Scrummage is under the GNU Public Licence, version 3, which provides no warranty for the software. Scrummage is Free, Open-Source, Software (FOSS).\n\nIf assistance is needed, the Scrummage project offers focused support in the form of sponsorship on the main GitHub page, the level of support depends on the level of sponsorship. Any issues with the product or this report, that are not unique to you or your organisation can also be raised as an issue on GitHub.")
                document.add_heading("Overview", 1)
                document.add_paragraph(f"This report highlights findings produced by the {Plugin_Name} plugin, that is part of the Scrummage project.")
                document.add_page_break()
                document.add_heading("Detailed Technical Findings", 1)
                document.add_paragraph("\n")
                print(5)

            document.add_heading(Title, 2)

            Document_Data = (
                ('Plugin', Plugin_Name),
                ('Domain', Domain),
                ('Link', Link),
                ('Created At', str(Date())),
                ('Result Type', Result_Type),
                ('Output Files', Output_File),
                ('Associated Task ID', str(Task_ID))
            )

            for name, data in Document_Data:
                document.add_paragraph(name, style=Finding_Style)
                document.add_paragraph(data)

            document.add_page_break()
            document.save(Output_File)
            logging.info(f"{Date()} - Common Library - Exported to DOCX file located at {str(Output_File)}.")
            return Output_File

        else:
            return None

    except Exception as e:
        logging.warning(f"{Date()} - Common Library - {str(e)}.")

def Load_Output(Object, Type):
    Standardised_Outputs: list = ["postgresql", "scumblr"]
    Defined_Outputs: dict = {
        "docx": ["use_docx"],
        "csv": ["use_csv"],
        "defectdojo": ["ssl", "api_key", "host", "user", "engagement_id", "product_id", "test_id", "user_id"],
        "rtir": ["ssl", "host", "port", "user", "password", "authenticator"],
        "jira": ["project_key", "address", "username", "password", "ticket_type"],
        "slack": ["token", "channel"],
        "elasticsearch": ["ssl", "host", "port", "index", "use_timestamp"],
        "email": ["smtp_server", "smtp_port", "from_address", "from_password", "to_address"]
    }

    if Type in Standardised_Outputs:
        return Object.Load_Configuration(Postgres_Database=True, Object=Type)

    elif Type in Defined_Outputs:
        return Object.Load_Configuration(Object=Type, Details_to_Load=Defined_Outputs[Type])

def Defect_Dojo_Output(Object, Title, Description):
    DD_Details = Load_Output(Object, "defectdojo")

    if DD_Details:

        try:

            if DD_Details[0]:
                Service: str = "https://"

            else:
                Service: str = "http://"

            Host = f"{Service}{DD_Details[2]}/api/v2/findings/"
            Impact: str = 'All Scrummage findings have the potential to cause significant damage to a business\' finances, efficiency and reputation. Therefore, findings should be investigated to assist in reducing this risk.'
            Mitigation: str = 'It is recommended that this issue be investigated further by the security team to determine whether or not further action needs to be taken.'
            DD_Finding_Payload = {
                "static_finding": True,
                "description": Description,
                "mitigation": Mitigation,
                "impact": Impact,
                "false_p": bool(),
                "date": str(datetime.datetime.now().strftime('%Y-%m-%d')),
                "severity": "Low",
                "numerical_severity": "1",
                "title": Title,
                "duplicate": bool(),
                "test": DD_Details[6],
                "found_by": [
                    1
                ],
                "verified": True,
                "line": int(),
                "active": True,
                "under_review": True,
                "dynamic_finding": True
            }
            Response = Request_Handler(url=Host, method="POST", json=DD_Finding_Payload, Optional_Headers={"Authorization": f"Token {DD_Details[1]}"}, Full_Response=True)
            
            if Response.status_code != 201:
                logging.info(f"{Date()} - Common Library - Failed to create DefectDojo finding. {str(e)}.")

            try:
                logging.info(f"{Date()} - Common Library - DefectDojo finding created.")

            except Exception as e:
                logging.info(f"{Date()} - Common Library - Failed to create DefectDojo finding. {str(e)}.")

        except (Exception, psycopg2.DatabaseError) as Error:
            logging.warning(Date() + str(Error))

def Main_Database_Insert(Object, Title, Plugin_Name, Domain, Link, Result_Type, Output_File, Task_ID, Screenshot_Path=False):
    Connection = Load_Output(Object, "postgresql")
    logging.info(f"{Date()} - Common Library - Loading Scrummage's Main Database configuration data.")

    if Connection:

        try:
            # Create connection cursor.
            Cursor = Connection.cursor()
            Cursor.execute("SELECT * FROM results WHERE link like %s", (Link,))
            Item_Already_in_Database = Cursor.fetchone()

            if Item_Already_in_Database is None:
                # Execute statement.

                if not Screenshot_Path:
                    Cursor.execute("INSERT INTO results (title, plugin, status, domain, link, created_at, updated_at, output_file, result_type, task_id) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)", (Title, Plugin_Name, "Open", Domain, Link, Date(), Date(), Output_File, Result_Type, Task_ID,))

                else:
                    Cursor.execute("INSERT INTO results (title, plugin, status, domain, link, created_at, updated_at, output_file, result_type, task_id, screenshot_url) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)", (Title, Plugin_Name, "Open", Domain, Link, Date(), Date(), Output_File, Result_Type, Task_ID, Screenshot_Path,))

            else:
                logging.info(f"{Date()} - Common Library - Entry already exists in the database. Skipping...")

        except (Exception, psycopg2.DatabaseError) as Error:
            logging.warning(Date() + str(Error))

        finally:

            if Connection is not None:
                Connection.commit()
                Connection.close()
                logging.info(f"{Date()} - Common Library - Result added to Scrummage database.")

def Scumblr_Main(Object, Link, Domain, Title):
    Connection = Load_Output(Object, "scumblr")

    if Connection:

        try:
            # Create connection cursor.
            Cursor = Connection.cursor()
            Cursor.execute("SELECT * FROM results WHERE url like %s", (Link,))

            if Cursor.fetchone() is None:
                # Execute statement.
                Cursor.execute("INSERT INTO results (title, url, created_at, updated_at, domain) VALUES(%s, %s, %s, %s, %s)", (Title, Link, Date(), Date(), Domain))

            else:
                logging.info(f"{Date()} - Common Library - Entry already exists in Scumblr database. Skipping...")

        except (Exception, psycopg2.DatabaseError) as Error:
            logging.warning(f"{Date()} - Common Library - " + Error)

        finally:

            if Connection is not None:
                Connection.commit()
                Connection.close()
                logging.info(f"{Date()} - Common Library - Result added to Scumblr database.")

def RTIR_Main(Object, Ticket_Subject, Ticket_Text):
    RTIR_Details = Load_Output(Object, "rtir")

    if RTIR_Details:

        try:

            if RTIR_Details[0]:
                Service: str = "https://"

            else:
                Service: str = "http://"

            Request_Data = f"content=id: ticket/new\nQueue: 1\nSubject: {Ticket_Subject}\nText: {Ticket_Text}"

            if RTIR_Details[5] != "cookie_based":
                logging.info(f"{Date()} - Common Library - No Authenticator specified, using the default which is cookie-based authentication.")

            RTIR_Response = Request_Handler(url=f"{Service}://{RTIR_Details[1]}:{RTIR_Details[2]}/REST/1.0/ticket/new?user={RTIR_Details[3]}&pass={RTIR_Details[4]}", method="POST", Data=Request_Data)

            if RTIR_Response.status_code == 200:
                logging.info(f"{Date()} - Common Library - New RTIR ticket created.")

            else:
                logging.warning(f"{Date()} - Common Library - Failed to create ticket in RTIR.")

        except Exception as e:
            logging.warning(f"{Date()} - Common Library - {str(e)}.")

def JIRA_Main(Object, Ticket_Summary, Ticket_Description):
    JIRA_Details = Load_Output(Object, "jira")

    if JIRA_Details:

        try:
            JIRA_Options={'server': JIRA_Details[1]}
            JIRA_Session=JIRA(options=JIRA_Options,basic_auth=(JIRA_Details[2], JIRA_Details[3]))
            JIRA_Session.create_issue(project={'key': JIRA_Details[0]}, summary=Ticket_Summary, description=Ticket_Description, issuetype={'name': JIRA_Details[4]})
            logging.info(f"{Date()} - Common Library - New JIRA ticket created.")

        except Exception as e:
            logging.warning(f"{Date()} - Common Library - {str(e)}.")

def Slack_Main(Object, Description):
    Slack_Details = Load_Output(Object, "slack")

    if Slack_Details:

        try:
            client = slack.WebClient(token=Slack_Details[0])
            client.chat_postMessage(channel=Slack_Details[1], text=Description)
            logging.info(f"{Date()} - Common Library - New Slack Notification created.")

        except Exception as e:
            logging.warning(f"{Date()} - Common Library - {str(e)}.")

def Elasticsearch_Main(Object, Title, Plugin_Name, Domain, Link, Result_Type, Output_File, Task_ID, Concat_Plugin_Name):
    Elasticsearch_Details = Load_Output(Object, "elasticsearch")

    if Elasticsearch_Details:

        try:

            if Elasticsearch_Details[4]:
                Timestamp = Date(Date_Only=True, Elastic=True)
                Index = Elasticsearch_Details[3] + "-" + Concat_Plugin_Name + "-" + Timestamp

            else:
                Index = Elasticsearch_Details[3] + "-" + Concat_Plugin_Name

            if Elasticsearch_Details[0]:
                Service: str = "https://"

            else:
                Service: str = "http://"

            URI = Service + Elasticsearch_Details[1] + ":" + str(Elasticsearch_Details[2]) + "/" + Index
            data = {"title": Title, "plugin": Plugin_Name, "domain": Domain, "link": Link, "output_file": Output_File, "result_type": Result_Type, "created_at": Date(), "associated_task_id": str(Task_ID)}
            data = JSON_Handler(data).Dump_JSON()
            resp = Request_Handler(url=URI, method="POST", Application_JSON_CT=True, Full_Response=True, json=data)

            if resp.status_code == 200:
                logging.info(f"{Date()} - Common Library - New result created in Elasticsearch, using the URI {URI}.")

            else:
                logging.info(f"{Date()} - Common Library - Failed to create result in Elasticsearch, using the URI {URI}.")

        except Exception as e:
            logging.warning(f"{Date()} - Common Library - Failed to create result in Elasticsearch. {str(e)}.")

def Email_Main(Object, Email_Subject, Email_Body):
    Email_Details = Load_Output(Object, "email")

    if Email_Details:

        try: # Send Email Alerts when called.
            server = smtplib.SMTP(Email_Details[0], Email_Details[1])
            server.ehlo()
            server.starttls()
            server.login(Email_Details[2], Email_Details[3])
            msg = MIMEMultipart()
            msg['From'] = Email_Details[2]
            msg['To'] = Email_Details[4]
            msg['Subject'] = Email_Subject
            msg.attach(MIMEText(Email_Body, 'plain'))
            text = msg.as_string()
            server.sendmail(Email_Details[2], Email_Details[4], text)
            server.quit()
            logging.info(f"{Date()} - Common Library - Email Sent.")

        except Exception as e:
            logging.warning(f"{Date()} - Common Library - Failed to send alert! Check email login settings. {str(e)}.")

def Set_Configuration_File():

    try:
        File_Dir = os.path.dirname(os.path.realpath('__file__'))
        return os.path.join(File_Dir, 'plugins/common/config/config.config')

    except Exception as e:
        logging.warning(f"DATE FUNCTION ERROR - Common Library - {str(e)}.")

def Date(Additional_Last_Days=0, Date_Only: bool = bool(), Elastic: bool = bool(), Full_Timestamp=False):

    try:

        def Date_Handler(Timestamp, Date_Only_Inner, Elastic_Inner):

            if Date_Only_Inner and not Elastic_Inner:
                return str(Timestamp.strftime('%Y-%m-%d'))

            elif Date_Only_Inner and Elastic_Inner:
                return str(Timestamp.strftime('%Y.%m.%d'))

            else:
                return str(Timestamp.strftime('%Y-%m-%d %H:%M:%S'))

        if Additional_Last_Days > 0:
            Additional_Last_Days_Range = Additional_Last_Days - 1
            Real_Dates: list = list()

            while Additional_Last_Days_Range < Additional_Last_Days and Additional_Last_Days_Range >= 0:
                Today = datetime.datetime.now()
                Day = datetime.timedelta(days=Additional_Last_Days_Range)
                Real_Date = Today - Day
                Real_Date = Date_Handler(Real_Date, Date_Only, Elastic)
                Real_Dates.append(Real_Date)
                Additional_Last_Days_Range -= 1

            return Real_Dates

        if Full_Timestamp:
            return datetime.datetime.now()

        else:
            return Date_Handler(datetime.datetime.now(), Date_Only, Elastic)

    except Exception as e:
        logging.warning(f"DATE FUNCTION ERROR - Common Library - {str(e)}.")

class JSON_Handler:

    def __init__(self, raw_data):
        self.json_data = raw_data

    def Is_JSON(self):

        try:
            json_object = json.loads(self.json_data)

        except ValueError:
            return False

        return json_object

    def To_JSON_Load(self):

        try:
            self.json_data = json.load(self.json_data)
            return self.json_data

        except Exception as e:
            logging.error(f"{Date()} - Common Library - {str(e)}.") 

    def To_JSON_Loads(self):

        try:
            self.json_data = json.loads(self.json_data)
            return self.json_data

        except Exception as e:
            logging.error(f"{Date()} - Common Library - {str(e)}.")   

    def Dump_JSON(self, Indentation=2, Sort=True):

        try:

            if Indentation > 0:
                self.json_data = json.dumps(self.json_data, indent=Indentation, sort_keys=Sort)

            else:
                self.json_data = json.dumps(self.json_data, sort_keys=Sort)

            return self.json_data

        except Exception as e:
            logging.error(f"{Date()} - Common Library - {str(e)}.")

def Request_Handler(User_Agent=True, Application_JSON_Accept: bool = bool(), Application_JSON_CT: bool = bool(), Application_Form_CT: bool = bool(), Accept_XML: bool = bool(), Accept_Language_EN_US: bool = bool(), Filter: bool = bool(), Risky_Plugin: bool = bool(), Full_Response: bool = bool(), Content_Response: bool = bool(), Host: str= str(), Data: dict = dict(), Params: dict = dict(), JSON_Data: dict = dict(), Optional_Headers: dict = dict(), Scrape_Regex_URL: str= str(), Proxies: dict = dict(), **kwargs):

    try:
        Headers: dict = dict()

        if not kwargs.get("method"):
            kwargs["method"] = "GET"

        if type(Data) in [dict, str]:
            Result = Configuration(Core=True).Load_Configuration(Object="proxy", Details_to_Load=["http", "https", "use_system_proxy"])

            if Result:
                Proxies: dict = dict()
                
                if Result[2]:
                    Proxies = urllib.request.getproxies()

                elif Result[0] != str() and Result[1] != str():
                    Proxies = {"http": Result[0], "https": Result[1]}
                
                if Proxies != dict():
                    kwargs["proxies"]: dict = Proxies

            if User_Agent:
                Headers['User-Agent'] = Current_User_Agent

            if Application_JSON_Accept:
                Headers['Accept']: str = 'application/json'
            
            if Application_JSON_CT:
                Headers['Content-Type']: str = 'application/json'

            if Accept_XML:
                Headers['Accept']: str = 'ext/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8'

            if Accept_Language_EN_US:
                Headers['Accept-Language']: str = 'en-US,en;q=0.5'

            if Application_Form_CT:
                Headers['Content-Type']: str = 'application/x-www-form-urlencoded'

            if type(Optional_Headers) == dict and len(Optional_Headers) > 0:

                for Header_Key, Header_Value in Optional_Headers.items():
                    Headers[Header_Key] = Header_Value

            kwargs["headers"]: dict = Headers
                    
            try:
                Response = requests.request(**kwargs)
                    
            except requests.exceptions.ConnectionError:
                return None

            if not Full_Response:

                if not Content_Response:
                    Response = Response.text

                else:
                    Response = Response.content

            Response_Dict: dict = dict()

            if Scrape_Regex_URL != str():
                Scrape_URLs: list = list()

                try:
                    Scrape_URLs_Raw = Regex_Handler(str(Response), Custom_Regex=Scrape_Regex_URL, Findall=True)

                    for Temp_URL_Extensions in Scrape_URLs_Raw:

                        if not Temp_URL_Extensions in Scrape_URLs:
                            Scrape_URLs.append(Temp_URL_Extensions)

                except Exception as e:
                    logging.warning(f"{Date()} - Common Library - Failed to regex URLs. {str(e)}.")

                Response_Dict["Regular"] = Response
                Response_Dict["Scraped"] = Scrape_URLs

            if Filter and str(Host) != str():
                To_Check = Response if not Full_Response else Response.text
                Filtered_Response = Response_Filter(To_Check, str(Host), Risky_Plugin=Risky_Plugin)

                if not Response_Dict.get("Regular"):
                    Response_Dict["Regular"] = Response

                Response_Dict["Filtered"] = Filtered_Response

            if Response_Dict != dict():
                return Response_Dict

            else:
                return Response

        else:
            logging.warning(f"{Date()} - Common Library - The data field needs to be in a dictionary or string format.")

    except Exception as e:
        logging.warning(f"{Date()} - Common Library - {str(e)}.")

def Response_Filter(Response, Host, Risky_Plugin=False):
    Risk_Level = Configuration(Core=True).Load_Configuration(Object="web_scraping", Details_to_Load=["risk_level", "automated_screenshots"])
    Risk_Level = Risk_Level[0]

    if (Risk_Level == 3) or (Risk_Level == 2 and not Risky_Plugin):
        Attributes = ["src", "href"]
        Quotes = ["\"", "\'"]
        Replace_Items = ["/", "./"]
        Replace_Strings = ["js", "assets", "polyfills", "main", "styles", "css", "jquery", "img", "images", "atom", "?", "static", "logo", "gui"]

        for Attribute in Attributes:

            for Quote in Quotes:
                Response = Response.replace(f"{Attribute} = {Quote}", f"{Attribute}={Quote}")

                if "https://" in Host:
                    Response = Response.replace(f"{Attribute}={Quote}//", f"{Attribute}={Quote}https://")

                else:
                    Response = Response.replace(f"{Attribute}={Quote}//", f"{Attribute}={Quote}http://")

                for Replace_Item in Replace_Items:
                    Response = Response.replace(f"{Attribute}={Quote}{Replace_Item}", f"{Attribute}={Quote}{Host}/")

                for Replace_String in Replace_Strings:
                    Response = Response.replace(f"{Attribute}={Quote}{Replace_String}", f"{Attribute}={Quote}{Host}/{Replace_String}")

                Response = Response.replace(f"{Attribute}={Quote}{Host}//", f"{Attribute}={Quote}{Host}/")
        
    return Response

def Regex_Handler(Query, Type: str= str(), Custom_Regex: str= str(), Findall: bool = bool(), Get_URL_Components=False):

    try:

        if Type != str():
            Predefined_Regex_Patterns = {"Phone": r"^\+\d+$", "Phone_Multi": r"^(\+)?\d+$", "Email": r"(^[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-\.]+$)", "Domain": r"([-a-zA-Z0-9@:%_\+~#=]{2,256}\.[a-z]{2,3})(\.[a-z]{2,3})?(\.[a-z]{2,3})?", "IP": r"[0-9]{1,3}\.[0-9]{1,3}\.[0-9]{1,3}\.[0-9]{1,3}", "URL": r"^(https?\:\/\/(www\.)?)?([a-zA-Z0-9@:%_\+\-~#=]{2,256})(\.[a-z]+)(\.[a-z]+)?(\.[a-z]+)?(\/)?$", "URL_Wild": r"(https?\:\/\/(www\.)?)?([-a-zA-Z0-9@:%_\+\-~#=]{2,256})(\.[a-z]{2,3})(\.[a-z]{2,3})?(\.[a-z]{2,3})?(\/)?.*", "MD5": r"([a-fA-F0-9]{32})\W", "SHA1": r"([a-fA-F0-9]{40})\W", "SHA256": r"([a-fA-F0-9]{64})\W", "Credentials": r"[\w\d\.\-\_]+\@[\w\.]+\:.*", "Cron": r"^([\d\/\*\-\,]+)\s([\d\/\*\-\,]+)\s([\d\/\*\-\,]+)\s([\d\/\*\-\,]+)\s([\d\/\*\-\,]+)$", "File_Date": r".+\/\d{4}\/\d{2}\/\d{2}\/.+", "Password_Special_Characters": r"[\@\_\-\!\#\$\%\^\&\*\(\)\~\`\<\>\]\[\}\{\|\:\;\'\"\/\?\.\,\+\=]+", "Company_Name": r".*[a-zA-Z].*", "Username": r"^[\s]+"}

            for Key, Value in Predefined_Regex_Patterns.items():

                if Type == Key:

                    if not Findall:
                        Regex = re.search(Value, Query)

                    else:
                        Regex = re.findall(Value, Query)

            if Regex and not Get_URL_Components:
                return Regex

            elif Regex and Type == "URL" and Get_URL_Components:
                URL_Prefix = Regex.group(1)
                URL_Body = Regex.group(3)

                if Regex.group(5) and Regex.group(6):
                    URL_Extension = Regex.group(4) + Regex.group(5) + Regex.group(6)

                elif Regex.group(5):
                    URL_Extension = Regex.group(4) + Regex.group(5)

                else:
                    URL_Extension = Regex.group(4)

                return {"Prefix": URL_Prefix, "Body": URL_Body, "Extension": URL_Extension}

            else:
                return None

        elif Custom_Regex != str():

            if not Findall:
                Regex = re.search(Custom_Regex, Query)

            else:
                Regex = re.findall(Custom_Regex, Query)

            if Regex:
                return Regex

            else:
                return None

        else:
            return None

    except Exception as e:
        logging.warning(f"{Date()} - Common Library - Failed to get check against a regex pattern. {str(e)}.")

def Filter(Segment_List, Start_Number, End_Number):

    try:

        def Dash_to_Numbers(Segment):
            List_of_Numbers: list = list()
            Segments = Segment.split("-")
            Iterator = int(Segments[0])

            while Iterator <= int(Segments[1]):
                List_of_Numbers.append(str(Iterator))
                Iterator += 1

            return List_of_Numbers

        Segment_List_Filtered: list = list()

        for Segment_Item in Segment_List:

            if "-" in Segment_Item:
                Segment_Item = Dash_to_Numbers(Segment_Item)
                Segment_List_Filtered.extend(Segment_Item)

            else:
                Segment_List_Filtered.append(Segment_Item)

        Non_Hardcoded_Segment_List: list = list()
        Updated_Segment_List: list = list()
        Iterator = 0
        Range_End = End_Number + 1
        Approved_Hours = list(range(Start_Number, Range_End))

        while Iterator < len(Segment_List_Filtered):
            Segment_Item = Segment_List_Filtered[Iterator]

            if Segment_Item != Segment_List_Filtered[-1] and "/" not in Segment_Item:
                Seg_Iter = 1
                First_Segment = Segment_Item
                Current_Segment = Segment_Item

                if (Iterator + Seg_Iter) < len(Segment_List_Filtered):
                    Current_Next_Value = Segment_List_Filtered[Iterator + Seg_Iter]
                    
                    while all(Seg.isnumeric() for Seg in [Current_Next_Value, Current_Segment]) and int(Current_Next_Value) in Approved_Hours and ((int(Current_Next_Value) - int(Current_Segment)) == 1):
                        Current_Segment = Current_Next_Value
                        Seg_Iter += 1
                        Curr_Iter = Iterator + Seg_Iter

                        if (Iterator + Seg_Iter) < len(Segment_List_Filtered):
                            Current_Next_Value = Segment_List_Filtered[Curr_Iter]

                        else:
                            break

                    if int(First_Segment) == End_Number:
                        Updated_Segment_List.append(First_Segment)

                    else:
                        Updated_Segment_List.append(First_Segment + "-" + Current_Segment)

                    Iterator += Seg_Iter

            elif "/" in Segment_Item:
                Non_Hardcoded_Segment_List.append(Segment_Item)
                Iterator += 1
            
            else:
                Iterator += 1

        if f"{str(Start_Number)}-{str(End_Number)}" in Updated_Segment_List:
            return ["*"]

        else:
            Updated_Segment_List.extend(Non_Hardcoded_Segment_List)
            return Updated_Segment_List

    except Exception as e:
        logging.warning(f"{Date()} - Common Library - Failed to verify and filter provided cron schedule. {str(e)}.")