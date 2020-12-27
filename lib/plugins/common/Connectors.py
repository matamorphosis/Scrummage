#!/usr/bin/env python3

import psycopg2, json, os, datetime, requests, slack, smtplib, csv, logging
from docx import Document
from jira.client import JIRA
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from defectdojo_api import defectdojo

File_Dir = os.path.dirname(os.path.realpath('__file__'))
Configuration_File = os.path.join(File_Dir, 'plugins/common/config/config.json')

def Date():
    return str(datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'))

def Load_Chrome_Configuration():
    logging.info(f"{Date()} Connectors Library - Loading RTIR configuration data.")

    try:
        with open(Configuration_File) as JSON_File:
            Configuration_Data = json.load(JSON_File)
            GC_Details = Configuration_Data['google-chrome']
            GC_App_Path = GC_Details['application-path']
            GC_Drv_Path = GC_Details['chromedriver-path']

            if GC_App_Path and GC_Drv_Path:
                return [GC_App_Path, GC_Drv_Path]

            else:
                return None

    except Exception as e:
        logging.warning(f"{Date()} Connectors Library - {str(e)}.")

def Load_CSV_Configuration():
    logging.info(f"{Date()} Connectors Library - Loading CSV configuration data.")

    try:
        with open(Configuration_File) as JSON_File:
            Configuration_Data = json.load(JSON_File)
            CSV_Details = Configuration_Data['csv']
            Use_CSV = CSV_Details['use-csv']

            if Use_CSV:
                return Use_CSV

            else:
                return None

    except Exception as e:
        logging.warning(f"{Date()} Connectors Library - {str(e)}.")

def Load_DOCX_Configuration():
    logging.info(f"{Date()} Connectors Library - Loading DOCX configuration data.")

    try:
        with open(Configuration_File) as JSON_File:
            Configuration_Data = json.load(JSON_File)
            DOCX_Details = Configuration_Data['docx-report']
            Use_DOCX = DOCX_Details['use-docx']

            if Use_DOCX:
                return Use_DOCX

            else:
                return None

    except Exception as e:
        logging.warning(f"{Date()} Connectors Library - {str(e)}.")

def Load_Defect_Dojo_Configuration():
    logging.info(f"{Date()} Connectors Library - Loading DefectDojo configuration data.")

    try:
        with open(Configuration_File) as JSON_File:
            Configuration_Data = json.load(JSON_File)
            DD_Details = Configuration_Data['defectdojo']
            DD_API_Key = DD_Details['api_key']
            DD_Host = DD_Details['host']
            DD_User = DD_Details['user']
            DD_Engagement_ID = DD_Details['engagement-id']
            DD_Product_ID = DD_Details['product-id']
            DD_Test_ID = DD_Details['test-id']
            DD_User_ID = DD_Details['user-id']

        if DD_API_Key and DD_Host and DD_User and DD_Engagement_ID and DD_Product_ID and DD_Test_ID and DD_User_ID:
            return [DD_API_Key, DD_Host, DD_User, DD_Engagement_ID, DD_Product_ID, DD_Test_ID, DD_User_ID]

        else:
            return None

    except Exception as e:
        logging.warning(f"{Date()} Connectors Library - {str(e)}.")

def Load_Email_Configuration():
    logging.info(f"{Date()} Connectors Library - Loading email configuration data.")

    try:
        with open(Configuration_File) as JSON_File:
            Configuration_Data = json.load(JSON_File)
            Email_Details = Configuration_Data['email']
            Email_SMTP_Server = Email_Details['smtp_server']
            Email_SMTP_Port = int(Email_Details['smtp_port'])
            Email_From_Address = Email_Details['from_address']
            Email_From_Password = Email_Details['from_password']
            Email_To_Address = Email_Details['to_address']

        if Email_SMTP_Server and Email_SMTP_Port and Email_From_Address and Email_From_Password and Email_To_Address:
            return [Email_SMTP_Server, Email_SMTP_Port, Email_From_Address, Email_From_Password, Email_To_Address]

        else:
            return None

    except Exception as e:
        logging.warning(f"{Date()} Connectors Library - {str(e)}.")

def Load_Elasticsearch_Configuration():
    logging.info(f"{Date()} Connectors Library - Loading Elasticsearch configuration data.")

    try:
        with open(Configuration_File) as JSON_File:
            Configuration_Data = json.load(JSON_File)
            Elasticsearch_Details = Configuration_Data['elasticsearch']
            Elasticsearch_Service = Elasticsearch_Details['service']
            Elasticsearch_Host = Elasticsearch_Details['host']
            Elasticsearch_Port = int(Elasticsearch_Details['port'])

        if Elasticsearch_Service and Elasticsearch_Host and Elasticsearch_Port:
            return [Elasticsearch_Service, Elasticsearch_Host, Elasticsearch_Port]

        else:
            return None

    except Exception as e:
        logging.warning(f"{Date()} Connectors Library - {str(e)}.")

def Load_Main_Database(Optional_File_Location=False):

    if Optional_File_Location:
        Current_Configuration_File = Optional_File_Location
    
    else:
        Current_Configuration_File = Configuration_File

    try:
        with open(Current_Configuration_File) as JSON_File:
            Configuration_Data = json.load(JSON_File)
            DB_Info = Configuration_Data['postgresql']
            DB_Host = DB_Info['host']
            DB_Port = str(int(DB_Info['port']))
            DB_Username = DB_Info['user']
            DB_Password = DB_Info['password']
            DB_Database = DB_Info['database']

    except:
        logging.warning(f"{Date()} Connectors Library - Failed to load configuration file.")

    try:
        DB_Connection = psycopg2.connect(user=DB_Username,
                                      password=DB_Password,
                                      host=DB_Host,
                                      port=DB_Port,
                                      database=DB_Database)

        if DB_Connection:
            return DB_Connection

        else:
            return None

    except:
        logging.warning(f"{Date()} Connectors Library - Failed to connect to database.")

def Load_JIRA_Configuration():
    logging.info(f"{Date()} Connectors Library - Loading JIRA configuration data.")

    try:

        with open(Configuration_File) as json_file:  
            Configuration_Data = json.load(json_file)
            JSON_Details = Configuration_Data['JIRA']
            JIRA_Project_Key = JSON_Details['project_key']
            JIRA_Address = JSON_Details['address']
            JIRA_Username = JSON_Details['username']
            JIRA_Password = JSON_Details['password']
            JIRA_Ticket_Type = JSON_Details['ticket_type']

            if JIRA_Project_Key and JIRA_Address and JIRA_Username and JIRA_Password and JIRA_Ticket_Type:
                return [JIRA_Project_Key, JIRA_Address, JIRA_Username, JIRA_Password, JIRA_Ticket_Type]

            else:
                return None

    except Exception as e:
        logging.warning(f"{Date()} Connectors Library - {str(e)}.")

def Load_Slack_Configuration():
    logging.info(f"{Date()} Connectors Library - Loading Slack configuration data.")

    try:

        with open(Configuration_File) as json_file:  
            Configuration_Data = json.load(json_file)
            JSON_Details = Configuration_Data['slack']
            Slack_Token = JSON_Details['token']
            Slack_Channel = JSON_Details['channel']

            if Slack_Token and Slack_Channel:
                return [Slack_Token, Slack_Channel]

            else:
                return None

    except Exception as e:
        logging.warning(f"{Date()} Connectors Library - {str(e)}.")

def Load_Scumblr_Configuration():
    logging.info(f"{Date()} Connectors Library - Loading Scumblr configuration data.")

    try:

        with open(Configuration_File) as JSON_File:
            Configuration_Data = json.load(JSON_File)
            PostgreSQL_Details = Configuration_Data['scumblr']
            PostgreSQL_Host = PostgreSQL_Details['host']
            PostgreSQL_Port = str(PostgreSQL_Details['port'])
            PostgreSQL_Database = PostgreSQL_Details['database']
            PostgreSQL_User = PostgreSQL_Details['user']
            PostgreSQL_Password = PostgreSQL_Details['password']

            if PostgreSQL_Host and PostgreSQL_Port and PostgreSQL_Database and PostgreSQL_User and PostgreSQL_Password:
                return [PostgreSQL_Host, PostgreSQL_Port, PostgreSQL_Database, PostgreSQL_User, PostgreSQL_Password]

            else:
                return None

    except Exception as e:
        logging.warning(f"{Date()} Connectors Library - {str(e)}.")

def Load_RTIR_Configuration():
    logging.info(f"{Date()} Connectors Library - Loading RTIR configuration data.")

    try:
        with open(Configuration_File) as JSON_File:
            Configuration_Data = json.load(JSON_File)
            RTIR_Details = Configuration_Data['rtir']
            RTIR_HTTP_Service  = RTIR_Details['service']
            RTIR_Host = RTIR_Details['host']
            RTIR_Port = str(RTIR_Details['port'])
            RTIR_User = RTIR_Details['user']
            RTIR_Password = RTIR_Details['password']
            RTIR_Authenticator = RTIR_Details['authenticator']

            if RTIR_HTTP_Service and RTIR_Host and RTIR_Port and RTIR_User and RTIR_Password and RTIR_Authenticator:
                return [RTIR_Host, RTIR_Port, RTIR_User, RTIR_Password, RTIR_HTTP_Service, RTIR_Authenticator]

            else:
                return None

    except Exception as e:
        logging.warning(f"{Date()} Connectors Library - {str(e)}.")

def CSV_Output(Title, Plugin_Name, Domain, Link, Result_Type, Output_File, Task_ID, Directory):

    try:
        Use_CSV = Load_CSV_Configuration()

        if Use_CSV:
            Headings = ["Title", "Plugin", "Domain", "Link", "Created At", "Output Files", "Result Type", "Task ID"]
            Data = [Title, Plugin_Name, Domain, Link, Date(), Output_File, Result_Type, str(Task_ID)]
            Complete_File = f"{File_Dir}/static/protected/output/{Directory}/{Plugin_Name}-Output.csv"

            if not os.path.exists(Complete_File):
                CSV_Output = csv.writer(open(Complete_File, 'w'))
                CSV_Output.writerow(Headings)
                CSV_Output.writerow(Data)
                logging.info(f"{Date()} Connectors Library - Created new CSV file located at {str(Complete_File)}.")

            else:
                CSV_Output = csv.writer(open(Complete_File, 'a'))
                CSV_Output.writerow(Data)
                logging.info(f"{Date()} Connectors Library - Updated existing CSV file located at {str(Complete_File)}.")

            return Complete_File

        else:
            return None

    except Exception as e:
        logging.warning(f"{Date()} Connectors Library - {str(e)}.")

def DOCX_Output(Title, Plugin_Name, Domain, Link, Result_Type, Output_File, Task_ID, Directory):

    try:
        Use_DOCX = Load_DOCX_Configuration()

        if Use_DOCX:
            Complete_File = f"{File_Dir}/static/protected/output/{Directory}/{Plugin_Name}-Output.docx"

            if os.path.exists(Complete_File):
                document = Document(Complete_File)

            else:
                from docx.shared import Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                document = Document()
                h1 = document.add_heading(f'Scrummage Finding Report for {Plugin_Name} Plugin', 0)
                h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                image = document.add_picture(f"{File_Dir}/static/images/search.png", width=Inches(2.00))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                document.add_page_break()

            Document_Data = (
                ('Plugin', Plugin_Name),
                ('Domain', Domain),
                ('Link', Link),
                ('Created At', str(Date())),
                ('Result Type', Result_Type),
                ('Output Files', Output_File),
                ('Associated Task ID', str(Task_ID))
            )

            table = document.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Title'
            hdr_cells[1].text = Title

            for name, data in Document_Data:
                row_cells = table.add_row().cells
                row_cells[0].text = name
                row_cells[1].text = data

            document.add_page_break()
            document.save(Complete_File)
            logging.info(f"{Date()} Connectors Library - Exported to DOCX file located at {str(Complete_File)}.")
            return Complete_File

        else:
            return None

    except Exception as e:
        logging.warning(f"{Date()} Connectors Library - {str(e)}.")


def Defect_Dojo_Output(Title, Description):
    DD_Details = Load_Defect_Dojo_Configuration()

    if DD_Details:

        try:
            Impact = 'All Scrummage findings have the potential to cause significant damage to a business\' finances, efficiency and reputation. Therefore, findings should be investigated to assist in reducing this risk.'
            Mitigation = 'It is recommended that this issue be investigated further by the security team to determine whether or not further action needs to be taken.'
            DD_Connection = defectdojo.DefectDojoAPI(DD_Details[1], DD_Details[0], DD_Details[2], debug=False)
            Finding = DD_Connection.create_finding(Title, Description, 'Low', '', str(datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S').strftime('%Y-%m-%d')), DD_Details[4], DD_Details[3], DD_Details[5], DD_Details[6], Impact, True, False, Mitigation)

            try:
                Finding = str(int(str(Finding)))
                logging.info(f"{Date()} Connectors Library - DefectDojo finding {Finding} created.")

            except:
                logging.info(f"{Date()} Connectors Library - Failed to create DefectDojo finding.")

        except (Exception, psycopg2.DatabaseError) as Error:
            logging.warning(Date() + str(Error))

def Main_Database_Insert(Title, Plugin_Name, Domain, Link, Result_Type, Output_File, Task_ID):
    Connection = Load_Main_Database()
    logging.info(f"{Date()} Connectors Library - Loading Scrummage's Main Database configuration data.")

    if Connection:

        try:
            # Create connection cursor.
            Cursor = Connection.cursor()
            Cursor.execute("SELECT * FROM results WHERE link like %s", (Link,))
            Item_Already_in_Database = Cursor.fetchone()

            if Item_Already_in_Database is None:
                # Execute statement.
                Cursor.execute("INSERT INTO results (title, plugin, status, domain, link, created_at, updated_at, output_file, result_type, task_id) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)", (Title, Plugin_Name, "Open", Domain, Link, Date(), Date(), Output_File, Result_Type, Task_ID,))

            else:
                logging.info(f"{Date()} Connectors Library - Entry already exists in the database. Skipping...")

        except (Exception, psycopg2.DatabaseError) as Error:
            logging.warning(Date() + str(Error))

        finally:

            if Connection is not None:
                Connection.commit()
                Connection.close()

def Scumblr_Main(Link, Domain, Title):
    Scumblr_Details = Load_Scumblr_Configuration()
    Connection = ""

    if Scumblr_Details:

        try:
            # Connect to the PostgreSQL server.
            Connection = psycopg2.connect(host=Scumblr_Details[0], port=Scumblr_Details[1], database=Scumblr_Details[2], user=Scumblr_Details[3], password=Scumblr_Details[4])

            # Create connection cursor.
            Cursor = Connection.cursor()
            Cursor.execute("SELECT * FROM results WHERE url like %s", (Link,))
            Item_Already_in_Database = Cursor.fetchone()

            if Item_Already_in_Database is None:
                # Execute statement.
                Cursor.execute("INSERT INTO results (title, url, created_at, updated_at, domain) VALUES(%s, %s, %s, %s, %s)", (Title, Link, Date(), Date(), Domain))

            else:
                logging.info(f"{Date()} Connectors Library - Entry already exists in Scumblr database. Skipping...")

        except (Exception, psycopg2.DatabaseError) as Error:
            logging.warning(f"{Date()} Connectors Library - " + Error)

        finally:

            if Connection is not None:
                Connection.commit()
                Connection.close()
                logging.info(f"{Date()} Connectors Library - Result added to Scumblr database.")
                logging.info(f"{Date()} Connectors Library - Database connection closed.")

def RTIR_Main(Ticket_Subject, Ticket_Text):
    RTIR_Details = Load_RTIR_Configuration()

    if RTIR_Details:

        try:
            Request_Data = f"content=id: ticket/new\nQueue: 1\nSubject: {Ticket_Subject}\nText: {Ticket_Text}"

            if RTIR_Details[5] != "cookie_based":
                logging.info(f"{Date()} Connectors Library - No Authenticator specified, using the default which is cookie-based authentication.")

            RTIR_Response = requests.post(f"{RTIR_Details[4]}://{RTIR_Details[0]}:{RTIR_Details[1]}/REST/1.0/ticket/new?user={RTIR_Details[2]}&pass={RTIR_Details[3]}", Request_Data)

            if RTIR_Response.status_code == 200:
                logging.info(f"{Date()} Connectors Library - New RTIR ticket created.")

            else:
                logging.warning(f"{Date()} Connectors Library - Failed to create ticket in RTIR.")

        except Exception as e:
            logging.warning(f"{Date()} Connectors Library - {str(e)}.")

def JIRA_Main(Ticket_Summary, Ticket_Description):
    JIRA_Details = Load_JIRA_Configuration()

    if JIRA_Details:

        try:
            JIRA_Options={'server': JIRA_Details[1]}
            JIRA_Session=JIRA(options=JIRA_Options,basic_auth=(JIRA_Details[2], JIRA_Details[3]))
            JIRA_Session.create_issue(project={'key': JIRA_Details[0]}, summary=Ticket_Summary, description=Ticket_Description, issuetype={'name': JIRA_Details[4]})
            logging.info(f"{Date()} Connectors Library - New JIRA ticket created.")

        except Exception as e:
            logging.warning(f"{Date()} Connectors Library - {str(e)}.")

def Slack_Main(Description):
    Slack_Details = Load_Slack_Configuration()

    if Slack_Details:

        try:
            client = slack.WebClient(token=Slack_Details[0])
            client.chat_postMessage(channel=Slack_Details[1], text=Description)
            logging.info(f"{Date()} Connectors Library - New Slack Notification created.")

        except Exception as e:
            logging.warning(f"{Date()} Connectors Library - {str(e)}.")

def Elasticsearch_Main(Title, Plugin_Name, Domain, Link, Result_Type, Output_File, Task_ID, Concat_Plugin_Name):
    Elasticsearch_Details = Load_Elasticsearch_Configuration()

    if Elasticsearch_Details:

        try:
            URI = Elasticsearch_Details[0] + Elasticsearch_Details[1] + ":" + str(Elasticsearch_Details[2]) + "/scrummage/result/" + Concat_Plugin_Name
            headers = {"Content-Type": "application/json"}
            data = {"title": Title, "plugin": Plugin_Name, "domain": Domain, "link": Link, "output_file": Output_File, "result_type": Result_Type, "created_at": Date(), "associated_task_id": str(Task_ID)}
            data = json.dumps(data)
            resp = requests.post(URI, data=data, headers=headers)

            if resp.status_code == 200:
                logging.info(f"{Date()} Connectors Library - New result created in Elasticsearch, using the URI " + URI + ".")

            else:
                logging.info(f"{Date()} Connectors Library - Failed to create result in Elasticsearch, using the URI " + URI + ".")

        except:
            logging.warning(f"{Date()} Connectors Library - Failed to create result in Elasticsearch.")

def Email_Main(Email_Subject, Email_Body):
    Email_Details = Load_Email_Configuration()

    if Email_Details:

        try: # Send Email Alerts when called.
            server = smtplib.SMTP(Email_Details[0], Email_Details[1])
            server.ehlo()
            server.starttls()
            server.login(Email_Details[2], Email_Details[3])
            msg = MIMEMultipart()
            msg['From'] = Email_Details[2]
            msg['To'] = Email_Details[4]
            msg['Subject'] = Email_Subject
            msg.attach(MIMEText(Email_Body, 'plain'))
            text = msg.as_string()
            server.sendmail(Email_Details[2], Email_Details[4], text)
            server.quit()
            logging.info(f"{Date()} Connectors Library - Email Sent.")

        except:
            logging.warning(f"{Date()} Connectors Library - Failed to send alert! Check email login settings.")